import os
from pathlib import Path
from typing import Optional


class DocumentLoader:
    """Load and extract text from various document formats."""

    SUPPORTED_TYPES = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".txt": "txt",
        ".md": "txt",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
        ".bmp": "image",
        ".tiff": "image",
    }

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        ext = Path(filename).suffix.lower()
        return ext in cls.SUPPORTED_TYPES

    @classmethod
    def get_type(cls, filename: str) -> Optional[str]:
        ext = Path(filename).suffix.lower()
        return cls.SUPPORTED_TYPES.get(ext)

    @staticmethod
    def load(file_path: str) -> str:
        """Load document and return extracted text.

        Args:
            file_path: Path to the document file

        Returns:
            Extracted text string
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext in (".pdf"):
            return DocumentLoader._load_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return DocumentLoader._load_docx(file_path)
        elif ext in (".txt", ".md"):
            return DocumentLoader._load_txt(file_path)
        elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff"):
            return DocumentLoader._load_image(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                text_parts.append(text.strip())
        doc.close()
        return "\n\n".join(text_parts)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n\n".join(text_parts)

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    @staticmethod
    def _load_image(file_path: str) -> str:
        from services.ocr_service import OCRService
        return OCRService.extract_text(file_path)
